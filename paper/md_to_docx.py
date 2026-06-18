"""Convierte el articulo Markdown a .docx con figuras y formulas.

Uso: py -3.12 paper/md_to_docx.py <md> <docx> <repo>
- Titulos, parrafos, negritas, listas y tablas -> estilos de Word.
- Imagenes  ![alt](ruta)  -> imagen incrustada.
- Formulas display  $$...$$  -> imagen renderizada con matplotlib (mathtext).
- Formulas en linea  $...$   -> texto con simbolos Unicode.
Punto de partida para volcar al template acmart.
"""
import os
import re
import sys
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FORM_DIR = None
_FORM_N = [0]

_INLINE = [
    (r"\\mathrm\{([^}]*)\}", r"\1"), (r"\\mathcal\{([^}]*)\}", r"\1"),
    (r"\\hat\{([^}]*)\}", r"\1̂"),
    (r"\\frac\{([^}]*)\}\{([^}]*)\}", r"\1/\2"),
    (r"\\,", " "), (r"\\;", " "), (r"\\cdot", "·"), (r"\\times", "×"),
    (r"\\le", "≤"), (r"\\ge", "≥"), (r"\\in", "∈"),
    (r"\\exists", "∃"), (r"\\wedge", "∧"), (r"\\sum", "Σ"),
    (r"\\pi", "π"), (r"\\sigma", "σ"), (r"\\log", "log"),
    (r"\\max", "max"), (r"\\min", "min"),
]


def _clean_inline(text):
    def repl(m):
        s = m.group(1)
        for pat, rep in _INLINE:
            s = re.sub(pat, rep, s)
        return s.replace("{", "").replace("}", "")
    return re.sub(r"\$([^$]+)\$", repl, text)


def _render_formula(tex):
    """Renderiza una formula LaTeX a PNG con mathtext. Devuelve la ruta o None."""
    _FORM_N[0] += 1
    path = os.path.join(_FORM_DIR, "eq_%02d.png" % _FORM_N[0])
    try:
        fig = plt.figure(figsize=(0.1, 0.1))
        fig.text(0.5, 0.5, "$%s$" % tex, ha="center", va="center", fontsize=17)
        fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.12)
        plt.close(fig)
        return path
    except Exception:
        plt.close("all")
        return None


def add_runs(paragraph, text):
    text = _clean_inline(text).replace("`", "")
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def md_to_docx(md_path, docx_path, repo):
    global _FORM_DIR
    base = os.path.dirname(os.path.abspath(md_path))
    _FORM_DIR = os.path.join(base, "assets", "formulas")
    os.makedirs(_FORM_DIR, exist_ok=True)

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line == "---":
            i += 1; continue

        # Formula display $$ ... $$
        if line.startswith("$$") and line.endswith("$$") and len(line) > 4:
            img = _render_formula(line.strip("$ ").strip())
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img:
                p.add_run().add_picture(img, width=Inches(4.6))
            else:
                p.add_run(_clean_inline(line.strip("$ "))).italic = True
            i += 1; continue

        # Imagen markdown ![alt](ruta)
        m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
        if m:
            rel = m.group(1)
            ip = os.path.normpath(os.path.join(base, rel))
            if os.path.exists(ip):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(ip, width=Inches(5.6))
            i += 1; continue

        # Tabla
        if line.startswith("|"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                buf.append(lines[i].strip()); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in buf]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        add_runs(t.cell(ri, ci).paragraphs[0], cell)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        elif re.match(r"^\d+\.\s", line):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            doc.add_paragraph().add_run(line.strip("*")).italic = True
        else:
            add_runs(doc.add_paragraph(), line)
        i += 1

    doc.save(docx_path)
    print("[OK] %s (%d formulas renderizadas)" % (docx_path, _FORM_N[0]))


if __name__ == "__main__":
    md_to_docx(sys.argv[1], sys.argv[2], sys.argv[3])
