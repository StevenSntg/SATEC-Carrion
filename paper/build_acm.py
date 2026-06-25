"""Genera el articulo en .docx con el estilo de revista de ACM (plantilla acmart).

A diferencia del generador anterior (estilos por defecto de Word + formulas como
imagenes PNG), este:

  * Usa la plantilla oficial **acmart** como base, heredando todos sus estilos
    profesionales (Titledocument, Authors, Abstract, Head1/2/3, DisplayFormula,
    FigureCaption, Bibentry, ...). Sin titulos azules: el estilo ACM es negro.
  * Renderiza las formulas como **ecuaciones nativas de Word (OMML)**, vectoriales
    y editables, via  latex -> latex2mathml -> MML2OMML.XSL (Office) -> OMML.
  * Numera automaticamente las ecuaciones (1..N) alineadas a la derecha.
  * Inserta las figuras con su pie (FigureCaption) y las tablas con su titulo
    (TableCaption), al ancho de la caja de texto.

Requisitos:
    pip install python-docx latex2mathml lxml pillow
    + Microsoft Office (usa MML2OMML.XSL para convertir MathML -> OMML nativo).

Uso:
    py -3.12 paper/build_acm.py            # md y docx por defecto
    py -3.12 paper/build_acm.py <md> <docx>
"""
import os
import re
import sys
import copy

from lxml import etree
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
TEMPLATE = r"C:\Users\Usuario\Desktop\IA\Sesion 9\acm_submission_template.docx"
XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

# Ancho de la caja de texto (8.5 - 1.0 - 1.42 pulgadas), para tabs y figuras.
TEXT_WIDTH_IN = 6.08

_XSLT = etree.XSLT(etree.parse(XSLT_PATH))

# ---------------------------------------------------------------------------
# Conversion LaTeX -> OMML nativo
# ---------------------------------------------------------------------------

# Preprocesado de comandos LaTeX que latex2mathml no interpreta de forma ideal.
# Los operadores con letras latinas (max, min) generan <mo> que Word renderiza en
# cursiva; convertirlos a \mathrm produce <mi mathvariant="normal"> que el XSLT
# marca como 'plain' (redonda), como corresponde a un operador. \mathrm{sen} ya
# da redonda, no se toca.
_TEX_FIX = [
    (r"\\!", ""),            # espacio negativo: irrelevante en OMML
    (r"\\,", " "), (r"\\;", " "), (r"\\:", " "), (r"\\quad", "  "),
    (r"\\qquad", "    "),
    (r"\\max", r"\\mathrm{max}"),
    (r"\\min", r"\\mathrm{min}"),
    (r"\\Big", ""), (r"\\big", ""), (r"\\Bigg", ""), (r"\\bigg", ""),
]


def _fix_tex(tex):
    for pat, rep in _TEX_FIX:
        tex = re.sub(pat, rep, tex)
    return tex


def latex_to_omml(tex):
    """Devuelve un elemento <m:oMath> (oxml) a partir de LaTeX, o None si falla."""
    import latex2mathml.converter as conv
    try:
        mathml = conv.convert(_fix_tex(tex))
        omml_root = _XSLT(etree.fromstring(mathml)).getroot()
        # Serializa con su propia declaracion xmlns:m (autocontenido) y reparsea
        # con el parser de python-docx para integrarlo al documento.
        return parse_xml(etree.tostring(omml_root).decode("utf-8"))
    except Exception as e:  # pragma: no cover - diagnostico
        sys.stderr.write("  [WARN] formula no convertida: %s (%s)\n"
                         % (tex[:50], e))
        return None


def strip_section_number(title):
    """Quita el numero manual de un encabezado ('2.1 Conjuntos' -> 'Conjuntos')."""
    return re.sub(r"^\d+(\.\d+)*\.?\s+", "", title)


def split_equations(tex):
    """Separa una linea con varias ecuaciones (P=.. \\qquad R=..) en piezas.

    Solo separa si CADA fragmento contiene '=', para no romper pares como
    'sen(...) , cos(...)' que constituyen una sola ecuacion.
    """
    parts = [p.strip() for p in re.split(r"\\qquad|\\quad", tex) if p.strip()]
    if len(parts) > 1 and all("=" in p for p in parts):
        return parts
    return [tex.strip()]


# ---------------------------------------------------------------------------
# Insercion de ecuaciones display (numeradas) e inline
# ---------------------------------------------------------------------------

def add_display_formula(doc, tex, number):
    """Anade una ecuacion display centrada con su numero (n) a la derecha."""
    p = doc.add_paragraph(style="DisplayFormula")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # los tabs hacen el centrado
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(TEXT_WIDTH_IN / 2), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Inches(TEXT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    omml = latex_to_omml(tex)
    if omml is not None:
        p._p.append(omml)
    else:
        p.add_run(tex)
    p.add_run("\t(%d)" % number)
    return p


def add_inline_omml(paragraph, tex):
    """Inserta una ecuacion inline (oMath) dentro de un parrafo existente."""
    omml = latex_to_omml(tex)
    if omml is not None:
        paragraph._p.append(omml)
        return True
    return False


# Limpieza minima para el fallback de inline a texto (si OMML fallara).
_INLINE_TXT = [
    (r"\\mathrm\{([^}]*)\}", r"\1"), (r"\\mathcal\{([^}]*)\}", r"\1"),
    (r"\\operatorname\{([^}]*)\}", r"\1"),
    (r"\\hat\{([^}]*)\}", r"\1̂"),
    (r"\\frac\{([^}]*)\}\{([^}]*)\}", r"\1/\2"),
    (r"\\times", "×"), (r"\\cdot", "·"), (r"\\to", "→"),
    (r"\\le", "≤"), (r"\\ge", "≥"), (r"\\in", "∈"),
    (r"\\wedge", "∧"), (r"\\pi", "π"), (r"\\sigma", "σ"),
    (r"\\,", " "), (r"_\{([^}]*)\}", r"_\1"), (r"\^\{([^}]*)\}", r"^\1"),
]


def _inline_fallback(tex):
    for pat, rep in _INLINE_TXT:
        tex = re.sub(pat, rep, tex)
    return tex.replace("{", "").replace("}", "")


# ---------------------------------------------------------------------------
# Runs con formato inline: **negrita**, *cursiva* y $formula$
# ---------------------------------------------------------------------------

_TOKEN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\$[^$]+\$|`[^`]+`)")


def add_runs(paragraph, text):
    """Agrega texto con negritas, cursivas, codigo y ecuaciones inline."""
    for part in _TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            paragraph.add_run(part[1:-1])
        elif part.startswith("$") and part.endswith("$"):
            tex = part[1:-1]
            if not add_inline_omml(paragraph, tex):
                paragraph.add_run(_inline_fallback(tex))
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Plantilla: limpiar el cuerpo conservando estilos y la seccion (sectPr)
# ---------------------------------------------------------------------------

def fresh_document():
    doc = Document(TEMPLATE)
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    sectPr_copy = copy.deepcopy(sectPr) if sectPr is not None else None
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl"), qn("w:sectPr")):
            body.remove(child)
    if sectPr_copy is not None:
        body.append(sectPr_copy)
    return doc


def img_width(path, target_in=4.5, max_h_in=3.2):
    """Ancho objetivo (centrado, < columna) con tope de altura para densidad.

    Las figuras apaisadas se insertan a ~4.9 in; las altas o cuadradas se
    reducen para no superar ~3.7 in de alto, evitando que dominen la pagina.
    """
    width = target_in
    try:
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        if (h / w) * width > max_h_in:
            width = max_h_in * (w / h)
    except Exception:
        pass
    return Inches(min(width, TEXT_WIDTH_IN))


def add_image(doc, path):
    p = doc.add_paragraph(style="Image")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=img_width(path))


# ---------------------------------------------------------------------------
# Tablas con estilo de revista (cabecera sombreada, sin bordes verticales)
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text, bold=False, align=None):
    cell.paragraphs[0].style = None
    cell.paragraphs[0].text = ""
    run_para = cell.paragraphs[0]
    if align is not None:
        run_para.alignment = align
    add_runs(run_para, text)
    for r in run_para.runs:
        r.font.size = Pt(8)
        if bold:
            r.font.bold = True


def add_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(t.cell(ri, ci), cell, bold=(ri == 0), align=align)
    _set_col_widths(t, rows)
    _style_table_borders(t)
    return t


def _set_col_widths(t, rows):
    """Ancho por columna proporcional a su contenido (1a columna mas ancha)."""
    ncols = len(rows[0])
    maxlen = [max(len(str(rows[r][c])) for r in range(len(rows)))
              for c in range(ncols)]
    # Primera columna segun su etiqueta mas larga, con tope; resto proporcional
    # al contenido y con un minimo para que no se parta el encabezado.
    w0 = min(1.7, max(1.1, TEXT_WIDTH_IN * maxlen[0] / max(sum(maxlen), 1)))
    rest_total = TEXT_WIDTH_IN - w0
    weights = maxlen[1:] or [1]
    data = [max(0.45, rest_total * w / sum(weights)) for w in weights]
    scale = rest_total / sum(data)
    widths = [Inches(w0)] + [Inches(w * scale) for w in data]
    tblPr = t._tbl.tblPr
    tblPr.append(parse_xml('<w:tblLayout %s w:type="fixed"/>' % _W))
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = widths[ci]


_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _style_table_borders(t):
    """Reglas horizontales tipo 'booktabs': arriba, bajo cabecera y abajo."""
    tblPr = t._tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="10" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="10" w:space="0" w:color="000000"/>'
        '<w:left w:val="nil"/><w:right w:val="nil"/>'
        '<w:insideH w:val="nil"/><w:insideV w:val="nil"/></w:tblBorders>' % _W)
    tblPr.append(borders)
    # Regla bajo la fila de cabecera (booktabs midrule) y sombreado suave.
    for cell in t.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            '<w:shd %s w:val="clear" w:color="auto" w:fill="ECECEC"/>' % _W))
        tcPr.append(parse_xml(
            '<w:tcBorders %s><w:bottom w:val="single" w:sz="6" w:space="0" '
            'w:color="000000"/></w:tcBorders>' % _W))


# ---------------------------------------------------------------------------
# Parser del articulo Markdown -> documento ACM
# ---------------------------------------------------------------------------

def build(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = fresh_document()
    eq_counter = [0]
    i = 0
    n = len(lines)
    prev_was_head = False
    in_refs = False

    def para(style):
        return doc.add_paragraph(style=style)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        # ---- Titulo principal ----
        if line.startswith("# "):
            p = para("Title_document")
            add_runs(p, line[2:].strip())
            prev_was_head = False
            i += 1
            continue

        # ---- Titulo corto: es metadato del running head, no va en el cuerpo ----
        if stripped.startswith("*Título corto") or stripped.startswith("*Titulo corto"):
            i += 1
            continue

        # ---- Autores y afiliaciones (front matter) ----
        m_auth = re.match(r"^\*\*([^*]+)\*\*,\s*(.+)$", stripped)
        if m_auth and not in_refs and eq_counter[0] == 0 and not prev_was_head:
            nombre, afil = m_auth.group(1).strip(), m_auth.group(2).strip()
            para("Authors").add_run(nombre)
            para("Affiliation").add_run(afil)
            i += 1
            continue

        # ---- Encabezados de seccion (acmart numera solo; quitamos el numero
        #      manual del Markdown para no duplicarlo) ----
        if line.startswith("### "):
            add_runs(para("Head2"), strip_section_number(line[4:].strip()))
            prev_was_head = True
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            low = title.lower()
            if low.startswith("abstract"):
                i += 1
                # primer parrafo siguiente -> Abstract (ingles)
                while i < n and not lines[i].strip():
                    i += 1
                ab = para("Abstract")
                add_runs(ab, lines[i].strip())
                i += 1
                continue
            if low.startswith("resumen"):
                i += 1
                while i < n and not lines[i].strip():
                    i += 1
                head = para("Abstract")
                r = head.add_run("Resumen.")
                r.bold = True
                r.italic = True
                head.add_run(" ")
                add_runs(head, lines[i].strip())
                i += 1
                continue
            if low.startswith("referencias") or low.startswith("references"):
                para("ReferenceHead").add_run("Referencias")
                in_refs = True
                prev_was_head = True
                i += 1
                continue
            if low.startswith("agradecimientos"):
                para("AckHead").add_run("Agradecimientos")
                prev_was_head = True
                i += 1
                # parrafo(s) de agradecimiento
                continue
            if low.startswith("disponibilidad"):
                para("AckHead").add_run(title)
                prev_was_head = True
                i += 1
                continue
            # seccion normal numerada
            add_runs(para("Head1"), strip_section_number(title))
            prev_was_head = True
            i += 1
            continue

        # ---- CCS Concepts ----
        if stripped.startswith("**CCS Concepts:**") or stripped.startswith("**CCS Concepts**"):
            txt = re.sub(r"^\*\*CCS Concepts:?\*\*\s*", "", stripped)
            p = para("CCSDescription")
            r = p.add_run("CCS Concepts: ")
            r.bold = True
            add_runs(p, txt)
            i += 1
            continue

        # ---- Keywords ----
        if stripped.startswith("**Keywords:**") or stripped.startswith("**Keywords**"):
            txt = re.sub(r"^\*\*Keywords:?\*\*\s*", "", stripped)
            p = para("KeyWords")
            r = p.add_run("Keywords: ")
            r.bold = True
            add_runs(p, txt)
            i += 1
            continue

        # ---- Formula display $$ ... $$ ----
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            inner = stripped[2:-2].strip()
            for piece in split_equations(inner):
                eq_counter[0] += 1
                add_display_formula(doc, piece, eq_counter[0])
            prev_was_head = False
            i += 1
            continue

        # ---- Imagen ![alt](ruta) ----
        m_img = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
        if m_img:
            rel = m_img.group(1)
            ip = os.path.normpath(os.path.join(os.path.dirname(md_path), rel))
            if os.path.exists(ip):
                add_image(doc, ip)
            prev_was_head = False
            i += 1
            continue

        # ---- Pie de figura  **Figura N.** ... ----
        if re.match(r"^\*\*Figura\s+\d+\.\*\*", stripped):
            p = para("FigureCaption")
            add_runs(p, stripped)
            i += 1
            continue

        # ---- Titulo de tabla  **Tabla N. ...** ----
        if re.match(r"^\*\*Tabla\s+\d+\.", stripped):
            p = para("TableCaption")
            add_runs(p, stripped.strip("* "))
            i += 1
            continue

        # ---- Tabla ----
        if stripped.startswith("|"):
            buf = []
            while i < n and lines[i].strip().startswith("|"):
                buf.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in buf]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                add_table(doc, rows)
            prev_was_head = False
            continue

        # ---- Referencias [n] ... (acmart numera solo; quitamos el [n] manual) ----
        if in_refs and re.match(r"^\[\d+\]", stripped):
            add_runs(para("Bib_entry"), re.sub(r"^\[\d+\]\s*", "", stripped))
            i += 1
            continue

        # ---- Linea en cursiva suelta (p. ej. nota) ----
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            para("Para").add_run(stripped.strip("*")).italic = True
            prev_was_head = False
            i += 1
            continue

        # ---- Parrafo normal (sin sangria de primera linea) ----
        style = "PostHeadPara" if prev_was_head else "Para"
        p = para(style)
        p.paragraph_format.first_line_indent = Pt(0)
        add_runs(p, stripped)
        prev_was_head = False
        i += 1

    doc.save(out_path)
    print("[OK] %s  (%d ecuaciones OMML)" % (out_path, eq_counter[0]))


if __name__ == "__main__":
    md = sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "SATEC_articulo_ACM.md")
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.join(HERE, "SATEC_articulo_ACM.docx")
    build(md, out)
